import io
from docx import Document
from app.documents.parser import DocumentParser, ParsedDocument, DocumentSection

class DocxParser(DocumentParser):
    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        doc = Document(io.BytesIO(content))
        
        sections = []
        current_title = "Document Start"
        current_content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            if para.style.name.startswith('Heading'):
                if current_content:
                    sections.append(DocumentSection(current_title, "\n".join(current_content), "text"))
                    current_content = []
                current_title = text
            else:
                current_content.append(text)
                
        if current_content:
            sections.append(DocumentSection(current_title, "\n".join(current_content), "text"))
            
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        
        return ParsedDocument(
            text=full_text,
            sections=sections,
            metadata={"source_type": "docx"}
        )
